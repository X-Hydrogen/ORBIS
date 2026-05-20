"""
Orbis Paper Generator — Main Generator

Generates a complete research paper from quantum chemistry results.
Outputs both LaTeX (.tex) and Word (.docx) formats.
"""

import json
import os
import re
import subprocess
import textwrap
from datetime import datetime
from pathlib import Path
from typing import Optional

from .figures import generate_all_figures

# ── LaTeX Article Template (ACS-style) ────────────────────────────

LATEX_PREAMBLE = r"""\documentclass[12pt,a4paper,titlepage]{article}

% ── Packages ──
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{amsmath,amssymb}
\usepackage{graphicx}
\usepackage[colorlinks=true,linkcolor=blue,citecolor=blue,urlcolor=blue]{hyperref}
\usepackage[margin=2.5cm]{geometry}
\usepackage{booktabs}
\usepackage{siunitx}
\usepackage[super,sort&compress,comma]{natbib}
\usepackage{float}
\usepackage{caption}
\usepackage{subcaption}
\usepackage{chemformula}
\usepackage{mhchem}
\usepackage{fancyhdr}
\usepackage{setspace}
\usepackage[version=4]{mhchem}

% ── Style ──
\onehalfspacing
\pagestyle{fancy}
\fancyhf{}
\fancyhead[L]{\small\itshape %%SHORT_TITLE%%}
\fancyhead[R]{\small\thepage}
\renewcommand{\headrulewidth}{0.4pt}

% ── Title/Author ──
\title{{\Large\bfseries %%TITLE%%}\\[0.3cm]
       {\large %%AUTHORS%%}\\[0.2cm]
       {\normalsize %%AFFILIATION%%}}

\author{}
\date{{\today}}

\begin{document}

\maketitle
\thispagestyle{fancy}

% ── Abstract ──
\begin{abstract}
\noindent
%%ABSTRACT%%
\end{abstract}

\newpage

% ══════════════════════════════════════════════════════════════════
% 1. Introduction
% ══════════════════════════════════════════════════════════════════
\section{Introduction}
\label{sec:introduction}

%%INTRODUCTION%%

% ══════════════════════════════════════════════════════════════════
% 2. Computational Methods
% ══════════════════════════════════════════════════════════════════
\section{Computational Methods}
\label{sec:methods}

%%METHODS%%

% ══════════════════════════════════════════════════════════════════
% 3. Results and Discussion
% ══════════════════════════════════════════════════════════════════
\section{Results and Discussion}
\label{sec:results}

%%RESULTS_DISCUSSION%%

% ══════════════════════════════════════════════════════════════════
% 4. Conclusions
% ══════════════════════════════════════════════════════════════════
\section{Conclusions}
\label{sec:conclusions}

%%CONCLUSIONS%%

% ══════════════════════════════════════════════════════════════════
% Acknowledgements
% ══════════════════════════════════════════════════════════════════
\section*{Acknowledgements}
\addcontentsline{toc}{section}{Acknowledgements}

%%ACKNOWLEDGEMENTS%%

% ══════════════════════════════════════════════════════════════════
% References
% ══════════════════════════════════════════════════════════════════
\begin{thebibliography}{99}

%%REFERENCES%%
\end{thebibliography}

\end{document}
"""

# ── Helper Functions ──────────────────────────────────────────────

def _sanitize_latex(text: str) -> str:
    """Escape special LaTeX characters."""
    replacements = {
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\^{}",
        "Å": r"\AA{}",
        "α": r"$\alpha$",
        "β": r"$\beta$",
        "γ": r"$\gamma$",
        "ω": r"$\omega$",
        "°": r"$^\circ$",
        "±": r"$\pm$",
        "×": r"$\times$",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def _wrap_text(text: str, width: int = 80) -> str:
    """Wrap text for LaTeX readability."""
    paragraphs = text.split("\n\n")
    wrapped = []
    for p in paragraphs:
        p = p.strip()
        if p:
            wrapped.append("\n".join(textwrap.wrap(p, width=width)))
    return "\n\n".join(wrapped)


def _detect_available_tools() -> dict:
    """Check which tools are available."""
    tools = {}
    # Check pandoc
    try:
        result = subprocess.run(["pandoc", "--version"], capture_output=True,
                                text=True, timeout=5)
        tools["pandoc"] = result.returncode == 0
    except Exception:
        tools["pandoc"] = False

    # Check pdflatex
    try:
        result = subprocess.run(["pdflatex", "--version"], capture_output=True,
                                text=True, timeout=5)
        tools["pdflatex"] = result.returncode == 0
    except Exception:
        tools["pdflatex"] = False

    # Check python-docx
    try:
        import docx
        tools["python-docx"] = True
    except ImportError:
        tools["python-docx"] = False

    return tools


# ── Reference Database ────────────────────────────────────────────

ORCA_REFERENCE = (
    r"F. Neese, F. Wennmohs, U. Becker, C. Riplinger, "
    r"J. Chem. Phys. \textbf{2020}, 152, 224108."
)

MULTIWFN_REFERENCE = (
    r"T. Lu, F. Chen, J. Comput. Chem. \textbf{2012}, 33, 580--592."
)

DEFAULT_REFERENCES = {
    "orca": ORCA_REFERENCE,
    "multiwfn": MULTIWFN_REFERENCE,
    "b3lyp": (
        r"A. D. Becke, J. Chem. Phys. \textbf{1993}, 98, 5648--5652; "
        r"C. Lee, W. Yang, R. G. Parr, Phys. Rev. B \textbf{1988}, 37, 785--789."
    ),
    "pbe0": (
        r"C. Adamo, V. Barone, J. Chem. Phys. \textbf{1999}, 110, 6158--6170."
    ),
    "def2": (
        r"F. Weigend, R. Ahlrichs, Phys. Chem. Chem. Phys. \textbf{2005}, 7, 3297--3305; "
        r"F. Weigend, Phys. Chem. Chem. Phys. \textbf{2006}, 8, 1057--1065."
    ),
    "d3bj": (
        r"S. Grimme, J. Antony, S. Ehrlich, H. Krieg, "
        r"J. Chem. Phys. \textbf{2010}, 132, 154104; "
        r"S. Grimme, S. Ehrlich, L. Goerigk, "
        r"J. Comput. Chem. \textbf{2011}, 32, 1456--1465."
    ),
    "rijcosx": (
        r"F. Neese, F. Wennmohs, A. Hansen, U. Becker, "
        r"Chem. Phys. \textbf{2009}, 356, 98--109; "
        r"R. Izs\'{a}k, F. Neese, J. Chem. Phys. \textbf{2011}, 135, 144105."
    ),
    "cp": (
        r"S. F. Boys, F. Bernardi, Mol. Phys. \textbf{1970}, 19, 553--566."
    ),
}


def _build_references(results: dict) -> str:
    """Build the reference list based on methods used."""
    refs = []
    ref_num = 0

    def add_ref(key, label):
        nonlocal ref_num
        ref_num += 1
        refs.append(
            rf"\bibitem{{{label}}} {DEFAULT_REFERENCES[key]}"
        )
        return ref_num

    methods = results.get("methods_used", {})

    # Always cite ORCA
    add_ref("orca", "orca")

    # Cite based on methods
    if methods.get("functional", "").upper() in ("B3LYP", "B3LYP-D3", "B3LYP-D3(BJ)"):
        add_ref("b3lyp", "b3lyp")
    if methods.get("functional", "").upper() in ("PBE0", "PBE0-D3", "PBE0-D3(BJ)"):
        add_ref("pbe0", "pbe0")
    if "def2" in methods.get("basis", ""):
        add_ref("def2", "def2")
    if "D3" in methods.get("dispersion", ""):
        add_ref("d3bj", "d3bj")
    if methods.get("ri_approx"):
        add_ref("rijcosx", "rijcosx")
    if methods.get("cp_correction"):
        add_ref("cp", "cp")
    if methods.get("multiwfn_used"):
        add_ref("multiwfn", "multiwfn")

    return "\n\n".join(refs)


# ══════════════════════════════════════════════════════════════════
# PaperGenerator Class
# ══════════════════════════════════════════════════════════════════

class PaperGenerator:
    """Generates complete research papers from quantum chemistry results."""

    def __init__(
        self,
        output_dir: str = None,
        title: str = "Computational Study of Molecular Properties",
        authors: str = "Hengyue Xu",
        affiliation: str = "Independent Researcher",
        short_title: str = "Computational Study",
    ):
        self.output_dir = Path(output_dir) if output_dir else Path.cwd()
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.title = title
        self.authors = authors
        self.affiliation = affiliation
        self.short_title = short_title

        # Subdirectories
        self.fig_dir = self.output_dir / "figures"
        self.tex_dir = self.output_dir / "tex"
        self.fig_dir.mkdir(exist_ok=True)
        self.tex_dir.mkdir(exist_ok=True)

        self.tools = _detect_available_tools()

    def generate(
        self,
        abstract: str,
        introduction: str,
        methods: str,
        results_discussion: str,
        conclusions: str,
        acknowledgements: str = "",
        results_data: dict = None,
    ) -> dict:
        """
        Generate the complete paper.

        Args:
            abstract: Abstract text
            introduction: Introduction section
            methods: Computational methods section
            results_discussion: Results and Discussion section
            conclusions: Conclusions section
            acknowledgements: Acknowledgements (optional)
            results_data: Dict with computed data for figures and references

        Returns:
            Dict with paths to generated files
        """
        results_data = results_data or {}

        # Generate figures
        figure_tex = ""
        try:
            figure_blocks = generate_all_figures(results_data, str(self.fig_dir))
            if figure_blocks:
                figure_tex = "\n\n".join(figure_blocks)
        except Exception as e:
            print(f"  ⚠️  Figure generation warning: {e}")
            figure_tex = "% Figures could not be generated automatically."

        # Insert figures into Results section (before the text)
        if figure_tex:
            results_discussion = figure_tex + "\n\n" + results_discussion

        # Build references
        references = _build_references(results_data)

        # Fill LaTeX template using simple replacement (avoids .format() conflicts with LaTeX braces)
        latex_content = LATEX_PREAMBLE
        replacements = {
            "%%TITLE%%": _sanitize_latex(self.title),
            "%%AUTHORS%%": _sanitize_latex(self.authors),
            "%%AFFILIATION%%": _sanitize_latex(self.affiliation),
            "%%SHORT_TITLE%%": _sanitize_latex(self.short_title),
            "%%ABSTRACT%%": _sanitize_latex(abstract),
            "%%INTRODUCTION%%": _sanitize_latex(introduction),
            "%%METHODS%%": _sanitize_latex(methods),
            "%%RESULTS_DISCUSSION%%": _sanitize_latex(results_discussion),
            "%%CONCLUSIONS%%": _sanitize_latex(conclusions),
            "%%ACKNOWLEDGEMENTS%%": _sanitize_latex(
                acknowledgements or
                "The authors thank the developers of ORCA and Multiwfn "
                "for making their software freely available for academic use."
            ),
            "%%REFERENCES%%": references,
        }
        for key, val in replacements.items():
            latex_content = latex_content.replace(key, str(val))

        # Write LaTeX file
        tex_path = self.tex_dir / "manuscript.tex"
        tex_path.write_text(latex_content)
        print(f"  ✅ LaTeX manuscript: {tex_path}")

        # Copy figures to tex dir for compilation
        for fig in self.fig_dir.glob("*.png"):
            import shutil
            shutil.copy(fig, self.tex_dir / fig.name)

        output = {
            "tex_path": str(tex_path),
            "fig_dir": str(self.fig_dir),
        }

        # Try to compile PDF
        if self.tools["pdflatex"]:
            try:
                pdf_path = self._compile_latex(tex_path)
                if pdf_path:
                    output["pdf_path"] = str(pdf_path)
                    print(f"  ✅ PDF compiled: {pdf_path}")
            except Exception as e:
                print(f"  ⚠️  PDF compilation failed: {e}")

        # Generate Word document
        if self.tools["pandoc"]:
            try:
                docx_path = self._tex_to_docx(tex_path)
                if docx_path:
                    output["docx_path"] = str(docx_path)
                    print(f"  ✅ Word document: {docx_path}")
            except Exception as e:
                print(f"  ⚠️  Word conversion failed: {e}")
        elif self.tools["python-docx"]:
            try:
                docx_path = self._generate_docx_python(
                    abstract, introduction, methods,
                    results_discussion, conclusions, acknowledgements,
                    results_data
                )
                if docx_path:
                    output["docx_path"] = str(docx_path)
                    print(f"  ✅ Word document: {docx_path}")
            except Exception as e:
                print(f"  ⚠️  Word conversion failed: {e}")
        else:
            print("  ⚠️  Neither pandoc nor python-docx available. "
                  "Install with: pip install python-docx")

        return output

    def _compile_latex(self, tex_path: Path) -> Optional[Path]:
        """Compile LaTeX to PDF."""
        workdir = tex_path.parent
        # Run twice for references
        for _ in range(2):
            result = subprocess.run(
                ["pdflatex", "-interaction=nonstopmode",
                 "-output-directory", str(workdir), tex_path.name],
                capture_output=True, text=True,
                cwd=str(workdir), timeout=60
            )
        pdf_path = workdir / "manuscript.pdf"
        if pdf_path.exists():
            return pdf_path
        return None

    def _tex_to_docx(self, tex_path: Path) -> Optional[Path]:
        """Convert LaTeX to DOCX using pandoc."""
        docx_path = self.output_dir / "manuscript.docx"
        result = subprocess.run(
            ["pandoc", str(tex_path), "-o", str(docx_path),
             "--from=latex", "--to=docx",
             "--resource-path=" + str(self.fig_dir)],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0 and docx_path.exists():
            return docx_path
        # Even on warnings, file may exist
        if docx_path.exists():
            return docx_path
        return None

    def _generate_docx_python(
        self,
        abstract: str,
        introduction: str,
        methods: str,
        results_discussion: str,
        conclusions: str,
        acknowledgements: str,
        results_data: dict,
    ) -> Optional[Path]:
        """Generate Word document using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            return None

        doc = Document()

        # Configure styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.5

        # Title
        title_para = doc.add_heading(self.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Authors
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(self.authors).bold = True

        # Affiliation
        affil_para = doc.add_paragraph()
        affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        affil_para.add_run(self.affiliation).italic = True

        doc.add_paragraph()  # blank line

        # Abstract
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(abstract)

        doc.add_page_break()

        # Introduction
        doc.add_heading("1. Introduction", level=1)
        for para in introduction.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Methods
        doc.add_heading("2. Computational Methods", level=1)
        for para in methods.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Results and Discussion
        doc.add_heading("3. Results and Discussion", level=1)
        for para in results_discussion.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Add figures
        for fig_path in sorted(self.fig_dir.glob("*.png")):
            try:
                doc.add_picture(str(fig_path), width=Inches(4.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

        # Conclusions
        doc.add_heading("4. Conclusions", level=1)
        for para in conclusions.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Acknowledgements
        if acknowledgements:
            doc.add_heading("Acknowledgements", level=1)
            doc.add_paragraph(acknowledgements)

        # References
        doc.add_heading("References", level=1)
        refs = _build_references(results_data)
        for ref in refs.split("\n\n"):
            if ref.strip():
                # Clean up LaTeX formatting
                clean_ref = re.sub(r"\\textbf\{([^}]+)\}", r"\1", ref)
                clean_ref = re.sub(r"\\textit\{([^}]+)\}", r"\1", clean_ref)
                clean_ref = re.sub(r"\\\w+\{([^}]+)\}", r"\1", clean_ref)
                clean_ref = clean_ref.replace(r"\"", "").replace("{", "").replace("}", "")
                doc.add_paragraph(clean_ref.strip())

        docx_path = self.output_dir / "manuscript.docx"
        doc.save(str(docx_path))
        return docx_path


# ══════════════════════════════════════════════════════════════════
# Convenience function
# ══════════════════════════════════════════════════════════════════

def generate_paper(
    results_data: dict,
    output_dir: str,
    title: str = None,
    authors: str = None,
) -> dict:
    """
    One-shot paper generation from Orbis agent results.

    Args:
        results_data: Complete results dictionary from Orbis agent with keys:
            - title, authors, affiliation
            - abstract, introduction, methods, results_discussion, conclusions
            - acknowledgements (optional)
            - Computational data for figures
        output_dir: Where to save the paper
        title: Override title
        authors: Override authors

    Returns:
        Dict with paths: tex_path, pdf_path (if compiled), docx_path (if generated)
    """
    gen = PaperGenerator(
        output_dir=output_dir,
        title=title or results_data.get("title", "Computational Study"),
        authors=authors or results_data.get("authors", "Hengyue Xu"),
        affiliation=results_data.get("affiliation", "Independent Researcher"),
        short_title=results_data.get("short_title",
                                     title or "Computational Study"),
    )

    return gen.generate(
        abstract=results_data.get("abstract", ""),
        introduction=results_data.get("introduction", ""),
        methods=results_data.get("methods", ""),
        results_discussion=results_data.get("results_discussion", ""),
        conclusions=results_data.get("conclusions", ""),
        acknowledgements=results_data.get("acknowledgements", ""),
        results_data=results_data,
    )


# ══════════════════════════════════════════════════════════════════
# CLI
# ══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Orbis Paper Generator"
    )
    parser.add_argument("--data", "-d", required=True,
                        help="JSON file with paper data")
    parser.add_argument("--output", "-o", default="./paper_output",
                        help="Output directory")
    args = parser.parse_args()

    data = json.loads(Path(args.data).read_text())
    result = generate_paper(data, args.output)
    print(f"\n📄 Paper generated in: {args.output}")
    for fmt, path in result.items():
        print(f"   {fmt}: {path}")
